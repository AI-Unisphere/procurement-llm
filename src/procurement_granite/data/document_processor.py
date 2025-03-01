"""Document processing utilities for the procurement-granite project."""

import os
from pathlib import Path
from typing import Dict, List, Optional, Union

import PyPDF2
import docx
from tqdm import tqdm

from procurement_granite.utils.config import get_data_path, load_config


class DocumentProcessor:
    """Process procurement documents from various formats into text."""
    
    def __init__(self, config_path: Optional[str] = None):
        """Initialize the document processor.
        
        Args:
            config_path: Path to the configuration file. If None, uses the default config.
        """
        self.config = load_config(config_path)
        self.supported_formats = self.config["data_processing"]["document_formats"]
    
    def process_document(self, file_path: Union[str, Path]) -> str:
        """Process a document file into text.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            str: The extracted text from the document.
            
        Raises:
            ValueError: If the file format is not supported.
            FileNotFoundError: If the file does not exist.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower().lstrip(".")
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: {self.supported_formats}")
        
        if file_extension == "pdf":
            return self._process_pdf(file_path)
        elif file_extension == "docx":
            return self._process_docx(file_path)
        elif file_extension == "txt":
            return self._process_txt(file_path)
        else:
            # This should not happen due to the check above, but just in case
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _process_pdf(self, file_path: Path) -> str:
        """Process a PDF file into text.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            str: The extracted text from the PDF.
        """
        text = ""
        
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
        
        return text
    
    def _process_docx(self, file_path: Path) -> str:
        """Process a DOCX file into text.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            str: The extracted text from the DOCX.
        """
        doc = docx.Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
        
        return text
    
    def _process_txt(self, file_path: Path) -> str:
        """Process a TXT file into text.
        
        Args:
            file_path: Path to the TXT file.
            
        Returns:
            str: The extracted text from the TXT file.
        """
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        
        return text
    
    def process_directory(self, directory_path: Union[str, Path], output_dir: Optional[Union[str, Path]] = None) -> Dict[str, str]:
        """Process all supported documents in a directory.
        
        Args:
            directory_path: Path to the directory containing documents.
            output_dir: Path to the directory to save processed text files. If None, files are not saved.
            
        Returns:
            Dict[str, str]: A dictionary mapping file names to extracted text.
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"Invalid directory path: {directory_path}")
        
        if output_dir is not None:
            output_dir = Path(output_dir)
            os.makedirs(output_dir, exist_ok=True)
        
        results = {}
        
        # Get all files with supported extensions
        files = []
        for ext in self.supported_formats:
            files.extend(list(directory_path.glob(f"*.{ext}")))
        
        for file_path in tqdm(files, desc="Processing documents"):
            try:
                text = self.process_document(file_path)
                results[file_path.name] = text
                
                if output_dir is not None:
                    output_path = output_dir / f"{file_path.stem}.txt"
                    with open(output_path, "w", encoding="utf-8") as f:
                        f.write(text)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
        
        return results


def process_raw_documents(output_dir: Optional[Union[str, Path]] = None) -> Dict[str, str]:
    """Process all raw documents in the data/raw directory.
    
    Args:
        output_dir: Path to the directory to save processed text files. 
                   If None, uses the default processed data directory.
    
    Returns:
        Dict[str, str]: A dictionary mapping file names to extracted text.
    """
    if output_dir is None:
        output_dir = get_data_path("processed")
    
    processor = DocumentProcessor()
    raw_data_path = get_data_path("raw")
    
    return processor.process_directory(raw_data_path, output_dir) 